# templates_app/views.py - Fichier complet avec exports fonctionnels
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth import login, authenticate
from django.contrib.auth.forms import UserCreationForm, AuthenticationForm
from django.contrib import messages
from django.http import HttpResponse, JsonResponse, Http404, FileResponse
from django.core.paginator import Paginator
from django.db.models import Q, Count
from django.utils import timezone
from django.views.decorators.cache import never_cache
from django.views.decorators.csrf import csrf_protect
from django.urls import reverse, NoReverseMatch
from django.template.loader import render_to_string
from .models import Template, Document, TemplateField, DocumentFieldValue, TemplateCategory
from .forms import TemplateForm, DocumentForm, TemplateFieldForm, TemplateCategoryForm
import re
import logging
import io
import os
from datetime import timedelta
from .forms import CustomUserCreationForm, CustomAuthenticationForm

# Imports pour les exports
# Imports pour PDF
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas

    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# Imports pour DOCX
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

# Alternative avec WeasyPrint - AVEC GESTION D'ERREUR
try:
    import weasyprint

    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False
    print("WeasyPrint non disponible - utilisation de ReportLab pour les PDFs")

# Configuration du logging pour debug
logger = logging.getLogger(__name__)


# ===============================
# VUE D'ACCUEIL
# ===============================

def home_view(request):
    """Page d'accueil avec statistiques et gestion du tutoriel"""
    context = {
        'total_templates': Template.objects.filter(is_public=True).count(),
        'total_users': Template.objects.values('created_by').distinct().count(),
        'total_documents': Document.objects.count(),
        'show_tutorial': not request.session.get('tutorial_completed', False)
    }
    return render(request, 'templates_app/home.html', context)


def complete_tutorial(request):
    """Marquer le tutoriel comme terminé"""
    if request.method == 'POST':
        request.session['tutorial_completed'] = True
        return JsonResponse({'success': True})
    return JsonResponse({'success': False})


# ===============================
# AUTHENTIFICATION
# ===============================

def signup_view(request):
    """Vue d'inscription"""
    if request.user.is_authenticated:
        return redirect('templates_app:template_list')

    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            username = form.cleaned_data.get('username')
            messages.success(request, f'Compte créé pour {username}! Vous pouvez maintenant vous connecter.')
            return redirect('templates_app:login')
    else:
        form = CustomUserCreationForm()

    return render(request, 'registration/signup.html', {'form': form})


def login_view(request):
    """Vue de connexion"""
    if request.user.is_authenticated:
        return redirect('templates_app:template_list')

    if request.method == 'POST':
        form = CustomAuthenticationForm(request, data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(username=username, password=password)
            if user is not None:
                login(request, user)
                messages.info(request, f'Vous êtes maintenant connecté en tant que {username}.')

                # Redirection après connexion
                next_page = request.GET.get('next')
                if next_page:
                    return redirect(next_page)
                else:
                    return redirect('templates_app:template_list')
        else:
            messages.error(request, 'Nom d\'utilisateur ou mot de passe incorrect.')
    else:
        form = CustomAuthenticationForm()

    return render(request, 'registration/login.html', {'form': form})


# ===============================
# GESTION DES TEMPLATES
# ===============================

@login_required
def template_list(request):
    """Liste des templates avec recherche et filtres"""
    # Récupération des paramètres de recherche
    search = request.GET.get('search', '').strip()
    category_filter = request.GET.get('category', '')
    my_templates = request.GET.get('my_templates', '') == 'on'

    # Construction de la requête
    templates = Template.objects.all()

    # Filtre par propriétaire
    if my_templates:
        templates = templates.filter(created_by=request.user)
    else:
        # Sinon, afficher les templates publics ou ceux de l'utilisateur
        templates = templates.filter(
            Q(created_by=request.user) | Q(is_public=True)
        )

    # Filtre par recherche
    if search:
        templates = templates.filter(
            Q(title__icontains=search) |
            Q(description__icontains=search) |
            Q(content__icontains=search)
        )

    # Filtre par catégorie
    if category_filter:
        templates = templates.filter(category_id=category_filter)

    # Tri par date de création (plus récent en premier)
    templates = templates.order_by('-created_at')

    # Pagination
    paginator = Paginator(templates, 12)  # 12 templates par page
    page_number = request.GET.get('page')
    templates = paginator.get_page(page_number)

    # Catégories pour le filtre
    categories = TemplateCategory.objects.all().order_by('name')

    context = {
        'templates': templates,
        'categories': categories,
        'search': search,
        'category_filter': category_filter,
        'my_templates': my_templates,
    }
    return render(request, 'templates_app/template_list.html', context)


@login_required
def template_detail(request, template_id):
    """Détail d'un template avec ses champs"""
    template = get_object_or_404(Template, id=template_id)

    # Vérifier les permissions
    if not template.is_public and template.created_by != request.user:
        messages.error(request, "Vous n'avez pas accès à ce template.")
        return redirect('templates_app:template_list')

    # Récupérer les champs du template
    fields = template.fields.all().order_by('order', 'field_name')

    # Compter les documents créés avec ce template
    documents_count = Document.objects.filter(template=template).count()
    user_documents_count = Document.objects.filter(
        template=template,
        created_by=request.user
    ).count() if request.user.is_authenticated else 0

    context = {
        'template': template,
        'fields': fields,
        'documents_count': documents_count,
        'user_documents_count': user_documents_count,
        'can_edit': request.user == template.created_by,
    }
    return render(request, 'templates_app/template_detail.html', context)


@login_required
def template_create(request):
    """Créer un nouveau template"""
    if request.method == 'POST':
        form = TemplateForm(request.POST)
        if form.is_valid():
            template = form.save(commit=False)
            template.created_by = request.user
            template.save()

            # Extraire et créer automatiquement les champs du contenu
            create_fields_from_content(template)

            messages.success(request, f'Template "{template.title}" créé avec succès!')
            return redirect('templates_app:template_detail', template_id=template.id)
    else:
        form = TemplateForm()

    categories = TemplateCategory.objects.all().order_by('name')
    context = {
        'form': form,
        'categories': categories,
        'action': 'Créer',
    }
    return render(request, 'templates_app/template_create.html', context)


@login_required
def template_edit(request, template_id):
    """Éditer un template existant"""
    template = get_object_or_404(Template, id=template_id, created_by=request.user)

    if request.method == 'POST':
        form = TemplateForm(request.POST, instance=template)
        if form.is_valid():
            form.save()

            # Mettre à jour les champs si le contenu a changé
            create_fields_from_content(template)

            messages.success(request, f'Template "{template.title}" modifié avec succès!')
            return redirect('templates_app:template_detail', template_id=template.id)
    else:
        form = TemplateForm(instance=template)

    categories = TemplateCategory.objects.all().order_by('name')
    context = {
        'form': form,
        'template': template,
        'categories': categories,
        'action': 'Modifier',
    }
    return render(request, 'templates_app/template_create.html', context)


# ===============================
# GESTION DES DOCUMENTS
# ===============================

@login_required
def document_create(request, template_id):
    """Créer un document basé sur un template"""
    template = get_object_or_404(Template, id=template_id)

    # Vérifier les permissions
    if not template.is_public and template.created_by != request.user:
        messages.error(request, "Vous n'avez pas accès à ce template.")
        return redirect('templates_app:template_list')

    # Récupérer les champs du template
    fields = template.fields.all().order_by('order', 'field_name')

    if request.method == 'POST':
        # Créer le document
        document_title = request.POST.get('document_title', f'Document basé sur {template.title}')
        document = Document.objects.create(
            title=document_title,
            template=template,
            created_by=request.user,
            is_completed=False
        )

        # Sauvegarder les valeurs des champs
        for field in fields:
            field_value = request.POST.get(f'field_{field.id}', '')
            if field_value:  # Ne sauvegarder que si il y a une valeur
                DocumentFieldValue.objects.create(
                    document=document,
                    template_field=field,
                    value=field_value
                )

        # Marquer comme complété si tous les champs requis sont remplis
        required_fields = fields.filter(is_required=True)
        filled_required = 0
        for field in required_fields:
            if request.POST.get(f'field_{field.id}'):
                filled_required += 1

        if filled_required == required_fields.count():
            document.is_completed = True
            document.save()

        messages.success(request, f'Document "{document.title}" créé avec succès!')
        return redirect('templates_app:document_detail', document_id=document.id)

    context = {
        'template': template,
        'fields': fields,
        'document': None,  # Pas de document existant pour la création
    }
    return render(request, 'templates_app/document_form.html', context)


@login_required
def document_detail(request, document_id):
    """Détail d'un document"""
    document = get_object_or_404(Document, id=document_id, created_by=request.user)

    # Vérifier que le document a bien un template
    if not document.template:
        messages.error(request, "Ce document n'est pas associé à un template valide.")
        return redirect('templates_app:document_list')

    template = document.template

    # Récupérer les valeurs des champs
    field_values = {}
    for field_value in document.field_values.all():
        field_name = field_value.template_field.field_name
        field_values[field_name] = field_value.value

    # Générer le contenu rendu
    rendered_content = document.template.content
    for field_name, field_value in field_values.items():
        placeholder = f'{{{{{field_name}}}}}'
        rendered_content = rendered_content.replace(placeholder, str(field_value))

    context = {
        'document': document,
        'template': template,
        'rendered_content': rendered_content,
        'field_values': field_values,
    }
    return render(request, 'templates_app/document_detail.html', context)


@login_required
def document_edit(request, document_id):
    """Éditer un document existant"""
    document = get_object_or_404(Document, id=document_id, created_by=request.user)
    template = document.template
    fields = template.fields.all().order_by('order', 'field_name')

    # Récupérer les valeurs existantes
    existing_values = {}
    for field_value in document.field_values.all():
        existing_values[field_value.template_field.id] = field_value.value

    if request.method == 'POST':
        # Mettre à jour le titre si fourni
        new_title = request.POST.get('document_title')
        if new_title and new_title != document.title:
            document.title = new_title
            document.save()

        # Supprimer les anciennes valeurs
        document.field_values.all().delete()

        # Sauvegarder les nouvelles valeurs
        for field in fields:
            field_value = request.POST.get(f'field_{field.id}', '')
            if field_value:
                DocumentFieldValue.objects.create(
                    document=document,
                    template_field=field,
                    value=field_value
                )

        # Mettre à jour le statut
        required_fields = fields.filter(is_required=True)
        filled_required = 0
        for field in required_fields:
            if request.POST.get(f'field_{field.id}'):
                filled_required += 1

        document.is_completed = filled_required == required_fields.count()
        document.save()

        messages.success(request, f'Document "{document.title}" modifié avec succès!')
        return redirect('templates_app:document_detail', document_id=document.id)

    context = {
        'document': document,
        'template': template,
        'fields': fields,
        'existing_values': existing_values,
    }
    return render(request, 'templates_app/document_edit.html', context)


@login_required
def document_list(request):
    """Liste des documents de l'utilisateur avec filtres"""
    # Paramètres de recherche et filtres
    search = request.GET.get('search', '').strip()
    selected_template = request.GET.get('template', '')
    selected_status = request.GET.get('status', '')

    # Requête de base
    documents = Document.objects.filter(created_by=request.user)
    all_documents = documents  # Pour les statistiques

    # Filtres
    if search:
        documents = documents.filter(
            Q(title__icontains=search) |
            Q(template__title__icontains=search)
        )

    if selected_template:
        documents = documents.filter(template_id=selected_template)

    # Correction: utiliser is_completed au lieu de status
    if selected_status == 'completed':
        documents = documents.filter(is_completed=True)
    elif selected_status == 'draft':
        documents = documents.filter(is_completed=False)

    # Tri par date de modification
    documents = documents.order_by('-updated_at')

    # Statistiques - Correction: utiliser is_completed
    completed_count = all_documents.filter(is_completed=True).count()
    draft_count = all_documents.filter(is_completed=False).count()

    # Documents créés cette semaine
    week_ago = timezone.now() - timedelta(days=7)
    recent_count = all_documents.filter(created_at__gte=week_ago).count()

    # Pagination
    paginator = Paginator(documents, 12)
    page_number = request.GET.get('page')
    documents = paginator.get_page(page_number)

    # Templates pour le filtre
    templates = Template.objects.filter(
        Q(created_by=request.user) | Q(is_public=True)
    ).order_by('title')

    context = {
        'documents': documents,
        'templates': templates,
        'search': search,
        'selected_template': selected_template,
        'selected_status': selected_status,
        'completed_count': completed_count,
        'draft_count': draft_count,
        'recent_count': recent_count,
    }
    return render(request, 'templates_app/document_list.html', context)


@login_required
def document_delete(request, document_id):
    """Supprimer un document avec confirmation"""
    document = get_object_or_404(Document, id=document_id, created_by=request.user)

    if request.method == 'POST':
        document_title = document.title
        document.delete()
        messages.success(request, f'Le document "{document_title}" a été supprimé avec succès.')
        return redirect('templates_app:document_list')

    context = {
        'document': document,
    }
    return render(request, 'templates_app/document_confirm_delete.html', context)


@login_required
def document_duplicate(request, document_id):
    """Dupliquer un document existant"""
    original_document = get_object_or_404(Document, id=document_id, created_by=request.user)

    # Créer une copie du document
    new_document = Document.objects.create(
        title=f"Copie de {original_document.title}",
        template=original_document.template,
        created_by=request.user,
        is_completed=False  # CORRECTION: utiliser is_completed au lieu de status
    )

    # Copier les valeurs des champs
    for field_value in original_document.field_values.all():
        DocumentFieldValue.objects.create(
            document=new_document,
            template_field=field_value.template_field,
            value=field_value.value
        )

    messages.success(request, f'Document dupliqué avec succès : "{new_document.title}"')
    return redirect('templates_app:document_detail', document_id=new_document.id)


# ===============================
# VUES D'EXPORT - FONCTIONNELLES
# ===============================

@login_required
def document_export_pdf(request, document_id):
    """Export PDF avec ReportLab ou WeasyPrint"""
    try:
        document = get_object_or_404(Document, id=document_id, created_by=request.user)

        if not REPORTLAB_AVAILABLE and not WEASYPRINT_AVAILABLE:
            messages.error(request,
                           "Les bibliothèques PDF ne sont pas installées. "
                           "Installez reportlab ou weasyprint avec: pip install reportlab")
            return redirect('templates_app:document_detail', document_id=document_id)

        # Vérifier que le document a un template
        if not document.template:
            messages.error(request, "Ce document n'a pas de template associé.")
            return redirect('templates_app:document_detail', document_id=document_id)

        # Récupérer les données du document
        template = document.template
        field_values = {}

        for field_value in document.field_values.all():
            field_name = field_value.template_field.field_name
            field_values[field_name] = field_value.value

        # Remplacer les placeholders
        rendered_content = template.content or ""
        for field_name, field_value in field_values.items():
            placeholder = f'{{{{{field_name}}}}}'
            rendered_content = rendered_content.replace(placeholder, str(field_value))

        if WEASYPRINT_AVAILABLE:
            return export_pdf_weasyprint(document, rendered_content)
        else:
            return export_pdf_reportlab(document, rendered_content)

    except Exception as e:
        # Log l'erreur pour debugging
        import traceback
        print(f"Erreur lors de l'export PDF: {str(e)}")
        print(traceback.format_exc())

        messages.error(request, f'Erreur lors de l\'export PDF : {str(e)}')
        return redirect('templates_app:document_detail', document_id=document_id)


@login_required
def document_export_docx(request, document_id):
    """Export DOCX avec python-docx"""
    document = get_object_or_404(Document, id=document_id, created_by=request.user)

    if not PYTHON_DOCX_AVAILABLE:
        messages.error(request,
                       "La bibliothèque python-docx n'est pas installée. "
                       "Installez-la avec: pip install python-docx")
        return redirect('templates_app:document_detail', document_id=document_id)

    try:
        # Récupérer les données du document
        template = document.template
        field_values = {}

        for field_value in document.field_values.all():
            field_name = field_value.template_field.field_name
            field_values[field_name] = field_value.value

        # Remplacer les placeholders
        rendered_content = template.content
        for field_name, field_value in field_values.items():
            placeholder = f'{{{{{field_name}}}}}'
            rendered_content = rendered_content.replace(placeholder, str(field_value))

        # Créer le document DOCX
        doc = DocxDocument()

        # Titre
        title = doc.add_heading(document.title, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Date
        date_para = doc.add_paragraph(f"Généré le {timezone.now().strftime('%d/%m/%Y à %H:%M')}")
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Contenu principal
        lines = rendered_content.split('\n')
        for line in lines:
            if line.strip():
                doc.add_paragraph(line.strip())

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph(f"Document généré par DocBuilder - {template.title}")
        footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Sauvegarder dans un buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Créer la réponse
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        filename = f"{document.title.replace(' ', '_')}.docx"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'

        return response

    except Exception as e:
        messages.error(request, f'Erreur lors de l\'export DOCX : {str(e)}')
        return redirect('templates_app:document_detail', document_id=document_id)


@login_required
def document_export_html(request, document_id):
    """Export HTML avec styles intégrés"""
    document = get_object_or_404(Document, id=document_id, created_by=request.user)

    try:
        # Récupérer les données du document
        template = document.template
        field_values = {}

        for field_value in document.field_values.all():
            field_name = field_value.template_field.field_name
            field_values[field_name] = field_value.value

        # Remplacer les placeholders
        rendered_content = template.content
        for field_name, field_value in field_values.items():
            placeholder = f'{{{{{field_name}}}}}'
            rendered_content = rendered_content.replace(placeholder, str(field_value))

        # Créer le HTML avec styles
        html_content = f"""<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{document.title}</title>
    <style>
        body {{
            font-family: 'Times New Roman', serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 2em auto;
            padding: 2em;
            color: #333;
        }}
        .document-header {{
            text-align: center;
            margin-bottom: 2em;
            padding-bottom: 1em;
            border-bottom: 2px solid #3498db;
        }}
        .document-header h1 {{
            color: #2c3e50;
            margin-bottom: 0.5em;
            font-size: 2em;
        }}
        .document-content {{
            margin: 2em 0;
            text-align: justify;
        }}
        .document-footer {{
            margin-top: 3em;
            padding-top: 1em;
            border-top: 1px solid #bdc3c7;
            font-size: 0.9em;
            color: #7f8c8d;
            text-align: center;
        }}
        h1, h2, h3 {{ color: #2c3e50; }}
        p {{ margin-bottom: 1em; }}

        @media print {{
            body {{ margin: 1cm; }}
            .no-print {{ display: none; }}
        }}
    </style>
</head>
<body>
    <div class="document-header">
        <h1>{document.title}</h1>
        <p><em>Généré le {timezone.now().strftime('%d/%m/%Y à %H:%M')}</em></p>
    </div>

    <div class="document-content">
        {rendered_content.replace(chr(10), '<br>')}
    </div>

    <div class="document-footer">
        <p><strong>Document généré par DocBuilder</strong></p>
        <p>Template source : {document.template.title}</p>
        <p class="no-print">
            <button onclick="window.print()" style="padding: 10px 20px; background: #3498db; color: white; border: none; border-radius: 5px; cursor: pointer;">
                Imprimer ce document
            </button>
        </p>
    </div>
</body>
</html>"""

        # Créer la réponse
        response = HttpResponse(html_content, content_type='text/html')
        filename = f"{document.title.replace(' ', '_')}.html"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'

        return response

    except Exception as e:
        messages.error(request, f'Erreur lors de l\'export HTML : {str(e)}')
        return redirect('templates_app:document_detail', document_id=document_id)


# ===============================
# FONCTIONS UTILITAIRES POUR EXPORTS
# ===============================

def export_pdf_weasyprint(document, content):
    """Export PDF avec WeasyPrint (recommandé pour le CSS)"""
    # Créer le HTML avec styles
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {{
                margin: 2cm;
                size: A4;
            }}
            body {{
                font-family: 'Times New Roman', serif;
                font-size: 12pt;
                line-height: 1.6;
                color: #333;
            }}
            h1, h2, h3 {{
                color: #2c3e50;
                margin-top: 1em;
                margin-bottom: 0.5em;
            }}
            h1 {{ font-size: 18pt; }}
            h2 {{ font-size: 16pt; }}
            h3 {{ font-size: 14pt; }}
            p {{ margin-bottom: 1em; }}
            .header {{
                text-align: center;
                margin-bottom: 2em;
                padding-bottom: 1em;
                border-bottom: 2px solid #3498db;
            }}
            .footer {{
                margin-top: 3em;
                padding-top: 1em;
                border-top: 1px solid #bdc3c7;
                font-size: 10pt;
                color: #7f8c8d;
            }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>{document.title}</h1>
            <p>Généré le {timezone.now().strftime('%d/%m/%Y à %H:%M')}</p>
        </div>

        <div class="content">
            {content.replace(chr(10), '<br>')}
        </div>

        <div class="footer">
            <p>Document généré par DocBuilder - {document.template.title}</p>
        </div>
    </body>
    </html>
    """

    # Générer le PDF
    pdf_file = weasyprint.HTML(string=html_content).write_pdf()

    # Créer la réponse HTTP
    response = HttpResponse(pdf_file, content_type='application/pdf')
    filename = f"{document.title.replace(' ', '_')}.pdf"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    return response


def export_pdf_reportlab(document, content):
    """Export PDF avec ReportLab (plus basique)"""
    buffer = io.BytesIO()

    # Créer le document PDF
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1 * inch)

    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontSize=18,
        spaceAfter=30,
        alignment=1  # Centré
    )

    normal_style = styles['Normal']
    normal_style.fontSize = 12
    normal_style.leading = 18

    # Contenu
    story = []

    # Titre
    title = Paragraph(document.title, title_style)
    story.append(title)

    # Date
    date_para = Paragraph(f"Généré le {timezone.now().strftime('%d/%m/%Y à %H:%M')}", styles['Normal'])
    story.append(date_para)
    story.append(Spacer(1, 20))

    # Contenu principal (diviser par lignes)
    lines = content.split('\n')
    for line in lines:
        if line.strip():
            para = Paragraph(line.strip(), normal_style)
            story.append(para)
        else:
            story.append(Spacer(1, 12))

    # Footer
    story.append(Spacer(1, 30))
    footer = Paragraph(f"Document généré par DocBuilder - {document.template.title}", styles['Normal'])
    story.append(footer)

    # Construire le PDF
    doc.build(story)

    # Réponse
    pdf = buffer.getvalue()
    buffer.close()

    response = HttpResponse(pdf, content_type='application/pdf')
    filename = f"{document.title.replace(' ', '_')}.pdf"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    return response


# ===============================
# GESTION DES CATÉGORIES
# ===============================

@login_required
def category_list(request):
    """Liste des catégories"""
    categories = TemplateCategory.objects.all().order_by('name')

    context = {
        'categories': categories,
    }
    return render(request, 'templates_app/category_list.html', context)


@login_required
def category_create(request):
    """Créer une nouvelle catégorie"""
    if request.method == 'POST':
        form = TemplateCategoryForm(request.POST)
        if form.is_valid():
            category = form.save()
            messages.success(request, f'Catégorie "{category.name}" créée avec succès!')
            return redirect('templates_app:category_list')
    else:
        form = TemplateCategoryForm()

    context = {
        'form': form,
        'action': 'Créer',
    }
    return render(request, 'templates_app/category_form.html', context)


@login_required
def category_edit(request, category_id):
    """Modifier une catégorie"""
    category = get_object_or_404(TemplateCategory, id=category_id)

    if request.method == 'POST':
        form = TemplateCategoryForm(request.POST, instance=category)
        if form.is_valid():
            form.save()
            messages.success(request, f'Catégorie "{category.name}" modifiée avec succès!')
            return redirect('templates_app:category_list')
    else:
        form = TemplateCategoryForm(instance=category)

    context = {
        'form': form,
        'category': category,
        'action': 'Modifier',
    }
    return render(request, 'templates_app/category_form.html', context)


@login_required
def category_delete(request, category_id):
    """Supprimer une catégorie"""
    category = get_object_or_404(TemplateCategory, id=category_id)

    if request.method == 'POST':
        category_name = category.name
        category.delete()
        messages.success(request, f'Catégorie "{category_name}" supprimée avec succès!')
        return redirect('templates_app:category_list')

    context = {
        'category': category,
    }
    return render(request, 'templates_app/category_confirm_delete.html', context)


# ===============================
# GESTION DES CHAMPS DE TEMPLATE
# ===============================

@login_required
def template_field_add(request, template_id):
    """Ajouter un champ à un template"""
    template = get_object_or_404(Template, id=template_id, created_by=request.user)

    if request.method == 'POST':
        form = TemplateFieldForm(request.POST)
        if form.is_valid():
            field = form.save(commit=False)
            field.template = template

            # Définir l'ordre du champ
            max_order = template.fields.aggregate(
                max_order=Count('order')
            )['max_order'] or 0
            field.order = max_order + 1

            field.save()
            messages.success(request, f'Champ "{field.field_name}" ajouté avec succès!')
            return redirect('templates_app:template_detail', template_id=template.id)
    else:
        form = TemplateFieldForm()

    context = {
        'form': form,
        'template': template,
        'action': 'Ajouter',
    }
    return render(request, 'templates_app/template_field_form.html', context)


@login_required
def template_field_edit(request, template_id, field_id):
    """Modifier un champ de template"""
    template = get_object_or_404(Template, id=template_id, created_by=request.user)
    field = get_object_or_404(TemplateField, id=field_id, template=template)

    if request.method == 'POST':
        form = TemplateFieldForm(request.POST, instance=field)
        if form.is_valid():
            form.save()
            messages.success(request, f'Champ "{field.field_label}" modifié avec succès!')
            return redirect('templates_app:template_detail', template_id=template.id)
    else:
        form = TemplateFieldForm(instance=field)

    context = {
        'form': form,
        'template': template,
        'field': field,
        'action': 'Modifier',
    }
    return render(request, 'templates_app/template_field_form.html', context)


@login_required
def template_field_delete(request, template_id, field_id):
    """Supprimer un champ de template"""
    template = get_object_or_404(Template, id=template_id, created_by=request.user)
    field = get_object_or_404(TemplateField, id=field_id, template=template)

    if request.method == 'POST':
        field_name = field.field_name
        field.delete()
        messages.success(request, f'Champ "{field_name}" supprimé avec succès.')
        return redirect('templates_app:template_detail', template_id=template.id)

    context = {
        'template': template,
        'field': field,
    }
    return render(request, 'templates_app/template_field_confirm_delete.html', context)


# ===============================
# VUES UTILITAIRES ET AJAX
# ===============================

@login_required
def template_preview(request, template_id):
    """Prévisualiser un template avec des données d'exemple (AJAX)"""
    template = get_object_or_404(Template, id=template_id)

    # Vérifier les permissions
    if not template.is_public and template.created_by != request.user:
        return JsonResponse({'error': 'Accès refusé'}, status=403)

    # Remplacer les placeholders par des exemples
    content = template.content
    fields = template.fields.all()

    for field in fields:
        placeholder = f'{{{{{field.field_name}}}}}'
        # CORRECTION : Utiliser getattr pour éviter l'erreur AttributeError
        if hasattr(field, 'default_value') and field.default_value:
            example_value = field.default_value
        else:
            # Générer des exemples basés sur le type de champ
            examples = {
                'text': f'[Exemple {field.field_name}]',
                'textarea': f'Texte d\'exemple pour {field.field_label}',
                'number': '123',
                'date': '2025-06-23',
                'email': 'exemple@email.com',
                'url': 'https://exemple.com',
                'select': 'Option 1' if field.get_options_list() else 'Sélection'
            }
            example_value = examples.get(field.field_type, f'[Exemple {field.field_name}]')

        content = content.replace(placeholder, example_value)

    return JsonResponse({
        'content': content,
        'title': template.title
    })


# 2. AJOUTER une nouvelle vue pour la page complète de preview
@login_required
def template_preview_page(request, template_id):
    """Vue complète pour la page de prévisualisation d'un template"""
    template = get_object_or_404(Template, id=template_id)

    # Vérifier les permissions
    if not template.is_public and template.created_by != request.user:
        messages.error(request, "Vous n'avez pas accès à ce template.")
        return redirect('templates_app:template_list')

    # Récupérer les champs
    fields = template.fields.all().order_by('order', 'field_name')

    # Générer le contenu avec des exemples
    rendered_content = template.content or ""
    preview_data = {}

    for field in fields:
        placeholder = f'{{{{{field.field_name}}}}}'

        # Générer une valeur d'exemple
        if hasattr(field, 'default_value') and field.default_value:
            example_value = field.default_value
        else:
            example_value = generate_sample_data(field.field_name)
            if not example_value:  # Si pas de données spécifiques
                examples = {
                    'text': f'Exemple de {field.field_label.lower()}',
                    'textarea': f'Texte d\'exemple pour {field.field_label}',
                    'number': '123',
                    'date': '2025-06-23',
                    'email': 'exemple@email.com',
                    'url': 'https://exemple.com',
                    'select': field.get_options_list()[0] if field.get_options_list() else 'Option 1'
                }
                example_value = examples.get(field.field_type, f'[Exemple {field.field_name}]')

        # Remplacer dans le contenu
        rendered_content = rendered_content.replace(placeholder, str(example_value))

        # Ajouter aux données de preview
        preview_data[field.field_name] = example_value

    context = {
        'template': template,
        'fields': fields,
        'rendered_content': rendered_content,
        'preview_data': preview_data,
    }
    return render(request, 'templates_app/template_preview.html', context)


@login_required
def document_preview(request, document_id):
    """Prévisualiser un document avec ses valeurs actuelles"""
    document = get_object_or_404(Document, id=document_id, created_by=request.user)

    # Récupérer les valeurs des champs
    field_values = {}
    for field_value in document.field_values.all():
        field_name = field_value.template_field.field_name
        field_values[field_name] = field_value.value

    # Générer le contenu rendu
    rendered_content = document.template.content
    for field_name, field_value in field_values.items():
        placeholder = f'{{{{{field_name}}}}}'
        rendered_content = rendered_content.replace(placeholder, str(field_value))

    return JsonResponse({
        'content': rendered_content,
        'title': document.title
    })


# ===============================
# VUES AVANCÉES
# ===============================

@login_required
def template_duplicate(request, template_id):
    """Dupliquer un template"""
    original_template = get_object_or_404(Template, id=template_id)

    # Vérifier les permissions
    if not original_template.is_public and original_template.created_by != request.user:
        messages.error(request, "Vous n'avez pas accès à ce template.")
        return redirect('templates_app:template_list')

    # Créer une copie du template
    new_template = Template.objects.create(
        title=f"Copie de {original_template.title}",
        description=original_template.description,
        content=original_template.content,
        category=original_template.category,
        created_by=request.user,
        is_public=False  # Les copies sont privées par défaut
    )

    # Copier les champs
    for field in original_template.fields.all():
        TemplateField.objects.create(
            template=new_template,
            field_name=field.field_name,
            field_type=field.field_type,
            order=field.order,
            is_required=field.is_required,
            field_options=field.field_options,
            placeholder_text=field.placeholder_text
        )

    messages.success(request, f'Template dupliqué avec succès : "{new_template.title}"')
    return redirect('templates_app:template_detail', template_id=new_template.id)


@login_required
def template_delete(request, template_id):
    """Supprimer un template avec confirmation"""
    template = get_object_or_404(Template, id=template_id, created_by=request.user)

    # Vérifier s'il y a des documents basés sur ce template
    document_count = Document.objects.filter(template=template).count()

    if request.method == 'POST':
        if document_count > 0:
            messages.error(request,
                           f'Impossible de supprimer le template "{template.title}" : '
                           f'{document_count} document(s) sont basés sur ce template.')
        else:
            template_title = template.title
            template.delete()
            messages.success(request, f'Template "{template_title}" supprimé avec succès.')
        return redirect('templates_app:template_list')

    context = {
        'template': template,
        'document_count': document_count,
    }
    return render(request, 'templates_app/template_confirm_delete.html', context)


# ===============================
# GESTION DES ERREURS
# ===============================

def handler404(request, exception):
    """Page d'erreur 404 personnalisée"""
    return render(request, 'templates_app/404.html', status=404)


def handler500(request):
    """Page d'erreur 500 personnalisée"""
    return render(request, 'templates_app/500.html', status=500)


# ===============================
# VUES DE RECHERCHE ET STATISTIQUES
# ===============================

@login_required
def search_global(request):
    """Recherche globale dans templates et documents"""
    query = request.GET.get('q', '').strip()

    if not query:
        return render(request, 'templates_app/search_results.html', {
            'query': query,
            'templates': [],
            'documents': [],
        })

    # Recherche dans les templates
    templates = Template.objects.filter(
        Q(created_by=request.user) | Q(is_public=True)
    ).filter(
        Q(title__icontains=query) |
        Q(description__icontains=query) |
        Q(content__icontains=query)
    ).order_by('-created_at')[:10]

    # Recherche dans les documents
    documents = Document.objects.filter(
        created_by=request.user
    ).filter(
        Q(title__icontains=query) |
        Q(template__title__icontains=query)
    ).order_by('-updated_at')[:10]

    context = {
        'query': query,
        'templates': templates,
        'documents': documents,
    }
    return render(request, 'templates_app/search_results.html', context)


@login_required
def dashboard_view(request):
    """Tableau de bord utilisateur avec statistiques détaillées - CORRIGÉ"""
    user = request.user

    # Récupérer les objets de l'utilisateur
    user_templates = Template.objects.filter(created_by=user)
    user_documents = Document.objects.filter(created_by=user)

    # Statistiques générales - CORRECTION: utiliser is_completed
    stats = {
        'total_templates': user_templates.count(),
        'public_templates': user_templates.filter(is_public=True).count(),
        'private_templates': user_templates.filter(is_public=False).count(),
        'total_documents': user_documents.count(),
        'completed_documents': user_documents.filter(is_completed=True).count(),  # CORRIGÉ
        'draft_documents': user_documents.filter(is_completed=False).count(),  # CORRIGÉ
    }

    # Activité récente (7 derniers jours)
    week_ago = timezone.now() - timedelta(days=7)
    stats['recent_templates'] = user_templates.filter(created_at__gte=week_ago).count()
    stats['recent_documents'] = user_documents.filter(created_at__gte=week_ago).count()

    # Activité récente détaillée
    recent_templates = user_templates.order_by('-created_at')[:5]
    recent_documents = user_documents.order_by('-updated_at')[:5]

    # Templates les plus utilisés
    template_usage = user_templates.annotate(
        doc_count=Count('documents')
    ).order_by('-doc_count')[:5]

    # Statistiques par catégorie
    categories_stats = []
    if TemplateCategory.objects.exists():
        for category in TemplateCategory.objects.all():
            category_templates = user_templates.filter(category=category)
            if category_templates.exists():
                categories_stats.append({
                    'category': category,
                    'count': category_templates.count(),
                    'color': category.color,
                })

    # Templates sans catégorie
    uncategorized_count = user_templates.filter(category__isnull=True).count()
    if uncategorized_count > 0:
        categories_stats.append({
            'category': {'name': 'Sans catégorie', 'color': '#6c757d'},
            'count': uncategorized_count,
            'color': '#6c757d',
        })

    # Progression mensuelle (derniers 6 mois)
    monthly_stats = []
    for i in range(6):
        month_start = timezone.now().replace(day=1) - timedelta(days=30 * i)
        month_end = month_start + timedelta(days=31)

        templates_count = user_templates.filter(
            created_at__gte=month_start,
            created_at__lt=month_end
        ).count()

        documents_count = user_documents.filter(
            created_at__gte=month_start,
            created_at__lt=month_end
        ).count()

        monthly_stats.insert(0, {
            'month': month_start.strftime('%B'),
            'templates': templates_count,
            'documents': documents_count,
        })

    # Suggestions d'amélioration
    suggestions = []
    if stats['total_templates'] == 0:
        suggestions.append({
            'type': 'info',
            'title': 'Créez votre premier template',
            'message': 'Commencez par créer un template pour vos documents récurrents.',
            'action_url': reverse('templates_app:template_create'),
            'action_text': 'Créer un template',
            'icon': 'fas fa-plus'
        })
    elif stats['public_templates'] == 0 and stats['total_templates'] > 0:
        suggestions.append({
            'type': 'success',
            'title': 'Partagez vos templates',
            'message': 'Rendez vos templates publics pour que d\'autres utilisateurs puissent en bénéficier.',
            'action_url': reverse('templates_app:template_list'),
            'action_text': 'Voir mes templates',
            'icon': 'fas fa-share'
        })

    if stats['draft_documents'] > stats['completed_documents'] and stats['draft_documents'] > 0:
        suggestions.append({
            'type': 'warning',
            'title': 'Finalisez vos brouillons',
            'message': f'Vous avez {stats["draft_documents"]} documents en brouillon qui attendent d\'être finalisés.',
            'action_url': reverse('templates_app:document_list') + '?status=draft',
            'action_text': 'Voir les brouillons',
            'icon': 'fas fa-clock'
        })

    context = {
        'stats': stats,
        'recent_templates': recent_templates,
        'recent_documents': recent_documents,
        'template_usage': template_usage,
        'categories_stats': categories_stats,
        'monthly_stats': monthly_stats,
        'suggestions': suggestions,
        'user': user,
    }

    return render(request, 'templates_app/dashboard.html', context)


# ===============================
# FONCTIONS UTILITAIRES
# ===============================

def extract_fields_from_content(content):
    """Extraire les champs dynamiques du contenu du template"""
    if not content:
        return []

    # Expression régulière pour trouver les placeholders {{field_name}}
    pattern = r'\{\{([a-zA-Z_][a-zA-Z0-9_]*)\}\}'
    matches = re.findall(pattern, content)

    # Supprimer les doublons et trier
    unique_fields = list(set(matches))
    unique_fields.sort()

    return unique_fields


def generate_sample_data(field_name):
    """Générer des données d'exemple pour un champ"""
    samples = {
        'nom_client': 'Jean Dupont',
        'nom_entreprise': 'ACME Corporation',
        'adresse_entreprise': '123 Rue de la Paix, 75001 Paris',
        'nom_representant': 'Marie Martin',
        'fonction_representant': 'Directrice des Ressources Humaines',
        'email': 'contact@example.com',
        'telephone': '01 23 45 67 89',
        'date_aujourd_hui': timezone.now().strftime('%d/%m/%Y'),
        'date_debut': '01/01/2024',
        'salaire': '3 500 €',
        'poste': 'Développeur Full-Stack',
        'montant_total': '1 250,00 €',
        'numero_reference': 'REF-2024-001',
        'ville': 'Paris',
        'date': timezone.now().strftime('%d/%m/%Y'),
    }

    # Essayer de deviner le type de données basé sur le nom du champ
    field_lower = field_name.lower()

    if 'email' in field_lower:
        return 'contact@example.com'
    elif 'date' in field_lower:
        return timezone.now().strftime('%d/%m/%Y')
    elif 'nom' in field_lower and 'client' in field_lower:
        return 'Jean Dupont'
    elif 'nom' in field_lower and 'entreprise' in field_lower:
        return 'ACME Corporation'
    elif 'nom' in field_lower and 'representant' in field_lower:
        return 'Marie Martin'
    elif 'fonction' in field_lower and 'representant' in field_lower:
        return 'Directrice des Ressources Humaines'
    elif 'adresse' in field_lower:
        return '123 Rue de la Paix, 75001 Paris'
    elif 'telephone' in field_lower or 'tel' in field_lower:
        return '01 23 45 67 89'
    elif 'montant' in field_lower or 'prix' in field_lower:
        return '1 250,00 €'
    elif 'numero' in field_lower or 'ref' in field_lower:
        return 'REF-2024-001'
    elif 'ville' in field_lower:
        return 'Paris'
    else:
        return samples.get(field_name, f'[Exemple pour {field_name.replace("_", " ")}]')


def create_fields_from_content(template):
    """Créer automatiquement les champs à partir du contenu du template"""
    if not template.content:
        return

    # Extraire les champs du contenu
    field_names = extract_fields_from_content(template.content)

    if not field_names:
        return

    # Créer les champs qui n'existent pas encore
    existing_fields = set(template.fields.values_list('field_name', flat=True))

    for i, field_name in enumerate(field_names):
        if field_name not in existing_fields:
            # Créer un libellé lisible à partir du nom du champ
            field_label = field_name.replace('_', ' ').title()

            # Déterminer le type de champ basé sur le nom
            field_type = 'text'  # Par défaut
            if 'email' in field_name.lower():
                field_type = 'email'
            elif 'date' in field_name.lower():
                field_type = 'date'
            elif 'description' in field_name.lower() or 'commentaire' in field_name.lower():
                field_type = 'textarea'
            elif 'montant' in field_name.lower() or 'prix' in field_name.lower() or 'salaire' in field_name.lower():
                field_type = 'number'

            # Déterminer si le champ est requis
            is_required = field_name.lower() in [
                'nom_client', 'nom_entreprise', 'email', 'date_debut', 'montant_total'
            ]

            TemplateField.objects.create(
                template=template,
                field_name=field_name,
                field_label=field_label,
                field_type=field_type,
                is_required=is_required,
                order=i * 10,
                placeholder_text=f'Entrez {field_label.lower()}'
            )


@login_required
def template_export(request, template_id):
    """Export de template en JSON ou autre format"""
    template = get_object_or_404(Template, id=template_id)

    # Vérifier les permissions
    if not template.is_public and template.created_by != request.user:
        messages.error(request, "Vous n'avez pas accès à ce template.")
        return redirect('templates_app:template_list')

    messages.info(request, "Export de template en cours de développement.")
    return redirect('templates_app:template_detail', template_id=template_id)


@login_required
def template_edit_fields(request, template_id):
    """Placeholder pour éditer les champs"""
    messages.info(request, "Fonction d'édition des champs en cours de développement.")
    return redirect('templates_app:template_detail', template_id=template_id)


@login_required
def template_add_field(request, template_id):
    """Placeholder pour ajouter un champ"""
    messages.info(request, "Fonction d'ajout de champ en cours de développement.")
    return redirect('templates_app:template_detail', template_id=template_id)


@login_required
def template_delete_field(request, template_id, field_id):
    """Supprimer un champ de template"""
    template = get_object_or_404(Template, id=template_id, created_by=request.user)
    field = get_object_or_404(TemplateField, id=field_id, template=template)

    if request.method == 'POST':
        field_name = field.field_name
        field.delete()
        messages.success(request, f'Champ "{field_name}" supprimé avec succès.')
        return redirect('templates_app:template_detail', template_id=template.id)

    context = {
        'template': template,
        'field': field,
    }
    return render(request, 'templates_app/template_field_confirm_delete.html', context)


@login_required
def dashboard_view(request):
    """Tableau de bord utilisateur avec statistiques détaillées"""
    user = request.user

    # Récupérer les objets de l'utilisateur
    user_templates = Template.objects.filter(created_by=user)
    user_documents = Document.objects.filter(created_by=user)

    # Statistiques générales
    stats = {
        'total_templates': user_templates.count(),
        'public_templates': user_templates.filter(is_public=True).count(),
        'private_templates': user_templates.filter(is_public=False).count(),
        'total_documents': user_documents.count(),
        'completed_documents': user_documents.filter(is_completed=True).count(),
        'draft_documents': user_documents.filter(is_completed=False).count(),
    }

    # Activité récente (7 derniers jours)
    week_ago = timezone.now() - timedelta(days=7)
    stats['recent_templates'] = user_templates.filter(created_at__gte=week_ago).count()
    stats['recent_documents'] = user_documents.filter(created_at__gte=week_ago).count()

    # Activité récente détaillée
    recent_templates = user_templates.order_by('-created_at')[:5]
    recent_documents = user_documents.order_by('-updated_at')[:5]

    # Templates les plus utilisés
    template_usage = user_templates.annotate(
        doc_count=Count('documents')
    ).order_by('-doc_count')[:5]

    # Statistiques par catégorie
    categories_stats = []
    if TemplateCategory.objects.exists():
        for category in TemplateCategory.objects.all():
            category_templates = user_templates.filter(category=category)
            if category_templates.exists():
                categories_stats.append({
                    'category': category,
                    'count': category_templates.count(),
                    'color': category.color,
                })

    # Templates sans catégorie
    uncategorized_count = user_templates.filter(category__isnull=True).count()
    if uncategorized_count > 0:
        categories_stats.append({
            'category': {'name': 'Sans catégorie', 'color': '#6c757d'},
            'count': uncategorized_count,
            'color': '#6c757d',
        })

    # Progression mensuelle (derniers 6 mois)
    monthly_stats = []
    for i in range(6):
        month_start = timezone.now().replace(day=1) - timedelta(days=30 * i)
        month_end = month_start + timedelta(days=31)

        templates_count = user_templates.filter(
            created_at__gte=month_start,
            created_at__lt=month_end
        ).count()

        documents_count = user_documents.filter(
            created_at__gte=month_start,
            created_at__lt=month_end
        ).count()

        monthly_stats.insert(0, {
            'month': month_start.strftime('%B'),
            'templates': templates_count,
            'documents': documents_count,
        })

    # Suggestions d'amélioration
    suggestions = []
    if stats['total_templates'] == 0:
        suggestions.append({
            'type': 'info',
            'title': 'Créez votre premier template',
            'message': 'Commencez par créer un template pour vos documents récurrents.',
            'action_url': reverse('templates_app:template_create'),
            'action_text': 'Créer un template',
            'icon': 'fas fa-plus'
        })
    elif stats['public_templates'] == 0 and stats['total_templates'] > 0:
        suggestions.append({
            'type': 'success',
            'title': 'Partagez vos templates',
            'message': 'Rendez vos templates publics pour que d\'autres utilisateurs puissent en bénéficier.',
            'action_url': reverse('templates_app:template_list'),
            'action_text': 'Voir mes templates',
            'icon': 'fas fa-share'
        })

    if stats['draft_documents'] > stats['completed_documents'] and stats['draft_documents'] > 0:
        suggestions.append({
            'type': 'warning',
            'title': 'Finalisez vos brouillons',
            'message': f'Vous avez {stats["draft_documents"]} documents en brouillon qui attendent d\'être finalisés.',
            'action_url': reverse('templates_app:document_list') + '?status=draft',
            'action_text': 'Voir les brouillons',
            'icon': 'fas fa-clock'
        })

    context = {
        'stats': stats,
        'recent_templates': recent_templates,
        'recent_documents': recent_documents,
        'template_usage': template_usage,
        'categories_stats': categories_stats,
        'monthly_stats': monthly_stats,
        'suggestions': suggestions,
        'user': user,
    }

    return render(request, 'templates_app/dashboard.html', context)


def complete_tutorial(request):
    """Marquer le tutoriel comme terminé (AJAX)"""
    if request.method == 'POST':
        try:
            # Marquer le tutoriel comme terminé dans la session
            request.session['tutorial_completed'] = True
            request.session.modified = True

            # Si l'utilisateur est connecté, on peut aussi le marquer dans le profil
            if request.user.is_authenticated:
                # Optionnel : enregistrer dans le profil utilisateur
                # request.user.profile.tutorial_completed = True
                # request.user.profile.save()
                pass

            return JsonResponse({
                'success': True,
                'message': 'Tutoriel marqué comme terminé !'
            })

        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': f'Erreur : {str(e)}'
            })

    # Si ce n'est pas une requête POST
    return JsonResponse({
        'success': False,
        'message': 'Méthode non autorisée'
    }, status=405)